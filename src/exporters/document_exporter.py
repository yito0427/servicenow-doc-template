"""
Document export functionality for multiple formats
"""
import io
import re
from pathlib import Path
from typing import Any, Dict, Optional

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import weasyprint
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class DocumentExporter:
    """Export documents to various formats"""
    
    def __init__(self):
        self.supported_formats = ['md', 'txt']
        
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
        
        if PDF_AVAILABLE:
            self.supported_formats.append('pdf')
    
    def export(self, content: str, format: str, output_path: Optional[Path] = None) -> bytes:
        """Export content to specified format"""
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}. Supported: {self.supported_formats}")
        
        if format == 'md':
            return content.encode('utf-8')
        elif format == 'txt':
            return self.export_to_text(content)
        elif format == 'docx':
            return self.export_to_word(content)
        elif format == 'pdf':
            return self.export_to_pdf(content)
        else:
            raise ValueError(f"Unknown format: {format}")
    
    def export_to_text(self, content: str) -> bytes:
        """Export to plain text"""
        # Remove markdown formatting
        text_content = re.sub(r'[#*`_~]', '', content)
        text_content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text_content)
        text_content = re.sub(r'^\|.*\|$', '', text_content, flags=re.MULTILINE)
        text_content = re.sub(r'^[-=]+$', '', text_content, flags=re.MULTILINE)
        text_content = re.sub(r'\n{3,}', '\n\n', text_content)
        
        return text_content.encode('utf-8')
    
    def export_to_word(self, content: str) -> bytes:
        """Export to Microsoft Word format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        doc = Document()
        
        # Parse markdown content
        lines = content.split('\n')
        current_table = []
        in_code_block = False
        code_content = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_table:
                    self._add_table_to_doc(doc, current_table)
                    current_table = []
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    self._add_code_block_to_doc(doc, '\n'.join(code_content))
                    code_content = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # Handle headers
            if line.startswith('#'):
                if current_table:
                    self._add_table_to_doc(doc, current_table)
                    current_table = []
                
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('# ').strip()
                
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(header_text)
                run.bold = True
                
                if level == 1:
                    run.font.size = Inches(0.2)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif level == 2:
                    run.font.size = Inches(0.18)
                elif level == 3:
                    run.font.size = Inches(0.16)
                
                continue
            
            # Handle tables
            if line.startswith('|') and line.endswith('|'):
                current_table.append(line)
                continue
            
            # Handle lists
            if line.startswith(('- ', '* ', '+ ')):
                if current_table:
                    self._add_table_to_doc(doc, current_table)
                    current_table = []
                
                list_text = line[2:].strip()
                paragraph = doc.add_paragraph(list_text, style='List Bullet')
                continue
            
            # Handle numbered lists
            if re.match(r'^\d+\.', line):
                if current_table:
                    self._add_table_to_doc(doc, current_table)
                    current_table = []
                
                list_text = re.sub(r'^\d+\.\s*', '', line)
                paragraph = doc.add_paragraph(list_text, style='List Number')
                continue
            
            # Regular paragraph
            if current_table:
                self._add_table_to_doc(doc, current_table)
                current_table = []
            
            if line:
                # Process markdown formatting
                processed_line = self._process_markdown_formatting(line)
                doc.add_paragraph(processed_line)
        
        # Add any remaining table
        if current_table:
            self._add_table_to_doc(doc, current_table)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def export_to_pdf(self, content: str) -> bytes:
        """Export to PDF format"""
        if not PDF_AVAILABLE:
            raise ImportError("weasyprint is required for PDF export")
        
        # Convert markdown to HTML
        html_content = self._markdown_to_html(content)
        
        # Create CSS styles
        css_styles = """
        <style>
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
            }
            h1 { color: #0066cc; border-bottom: 2px solid #0066cc; padding-bottom: 10px; }
            h2 { color: #0066cc; border-bottom: 1px solid #ccc; padding-bottom: 5px; }
            h3 { color: #0066cc; }
            h4, h5, h6 { color: #666; }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            th {
                background-color: #f5f5f5;
                font-weight: bold;
            }
            code {
                background-color: #f4f4f4;
                padding: 2px 4px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
            }
            pre {
                background-color: #f4f4f4;
                padding: 10px;
                border-radius: 5px;
                overflow-x: auto;
                font-family: 'Courier New', monospace;
            }
            ul, ol { margin: 10px 0; padding-left: 30px; }
            li { margin: 5px 0; }
            .page-break { page-break-before: always; }
        </style>
        """
        
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            {css_styles}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        pdf_buffer = io.BytesIO()
        weasyprint.HTML(string=full_html).write_pdf(pdf_buffer)
        return pdf_buffer.getvalue()
    
    def _process_markdown_formatting(self, text: str) -> str:
        """Process basic markdown formatting for Word"""
        # Remove markdown syntax for now (could be enhanced)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)        # Code
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links
        return text
    
    def _add_table_to_doc(self, doc: Any, table_lines: list):
        """Add table to Word document"""
        if len(table_lines) < 2:
            return
        
        # Parse table data
        rows = []
        for line in table_lines:
            if '---' in line or '===' in line:
                continue  # Skip separator lines
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_data
                    if i == 0:  # Header row
                        table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    
    def _add_code_block_to_doc(self, doc: Any, code_content: str):
        """Add code block to Word document"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code_content)
        run.font.name = 'Courier New'
        paragraph.style = 'No Spacing'
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML"""
        html = content
        
        # Headers
        html = re.sub(r'^# (.*$)', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.*$)', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^### (.*$)', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^#### (.*$)', r'<h4>\1</h4>', html, flags=re.MULTILINE)
        html = re.sub(r'^##### (.*$)', r'<h5>\1</h5>', html, flags=re.MULTILINE)
        html = re.sub(r'^###### (.*$)', r'<h6>\1</h6>', html, flags=re.MULTILINE)
        
        # Bold and italic
        html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html)
        
        # Code
        html = re.sub(r'`(.*?)`', r'<code>\1</code>', html)
        
        # Code blocks
        html = re.sub(r'```(.*?)```', r'<pre>\1</pre>', html, flags=re.DOTALL)
        
        # Links
        html = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', html)
        
        # Lists
        lines = html.split('\n')
        result_lines = []
        in_ul = False
        in_ol = False
        
        for line in lines:
            if re.match(r'^\s*[-*+]\s', line):
                if not in_ul:
                    result_lines.append('<ul>')
                    in_ul = True
                if in_ol:
                    result_lines.append('</ol>')
                    in_ol = False
                
                item_text = re.sub(r'^\s*[-*+]\s', '', line)
                result_lines.append(f'<li>{item_text}</li>')
            elif re.match(r'^\s*\d+\.\s', line):
                if not in_ol:
                    result_lines.append('<ol>')
                    in_ol = True
                if in_ul:
                    result_lines.append('</ul>')
                    in_ul = False
                
                item_text = re.sub(r'^\s*\d+\.\s', '', line)
                result_lines.append(f'<li>{item_text}</li>')
            else:
                if in_ul:
                    result_lines.append('</ul>')
                    in_ul = False
                if in_ol:
                    result_lines.append('</ol>')
                    in_ol = False
                
                if line.strip():
                    result_lines.append(f'<p>{line}</p>')
                else:
                    result_lines.append('<br>')
        
        if in_ul:
            result_lines.append('</ul>')
        if in_ol:
            result_lines.append('</ol>')
        
        html = '\n'.join(result_lines)
        
        # Tables
        html = self._convert_tables_to_html(html)
        
        return html
    
    def _convert_tables_to_html(self, content: str) -> str:
        """Convert markdown tables to HTML"""
        lines = content.split('\n')
        result = []
        in_table = False
        table_rows = []
        
        for line in lines:
            if line.strip().startswith('|') and line.strip().endswith('|'):
                if '---' in line or '===' in line:
                    continue  # Skip separator lines
                
                if not in_table:
                    in_table = True
                    table_rows = []
                
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                table_rows.append(cells)
            else:
                if in_table:
                    # End of table, convert to HTML
                    result.append(self._table_rows_to_html(table_rows))
                    table_rows = []
                    in_table = False
                
                result.append(line)
        
        if in_table and table_rows:
            result.append(self._table_rows_to_html(table_rows))
        
        return '\n'.join(result)
    
    def _table_rows_to_html(self, rows: list) -> str:
        """Convert table rows to HTML table"""
        if not rows:
            return ''
        
        html = ['<table>']
        
        # Header row
        if rows:
            html.append('<thead><tr>')
            for cell in rows[0]:
                html.append(f'<th>{cell}</th>')
            html.append('</tr></thead>')
        
        # Body rows
        if len(rows) > 1:
            html.append('<tbody>')
            for row in rows[1:]:
                html.append('<tr>')
                for cell in row:
                    html.append(f'<td>{cell}</td>')
                html.append('</tr>')
            html.append('</tbody>')
        
        html.append('</table>')
        return '\n'.join(html)